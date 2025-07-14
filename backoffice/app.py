from flask import Flask, request, jsonify
from flask_cors import CORS
import psycopg2
import docx
import logging
from fuzzywuzzy import fuzz
import faiss
from sentence_transformers import SentenceTransformer
import numpy as np
import pickle
import os
import traceback
import re
from unidecode import unidecode

# ---------- LOGGING ----------
logging.basicConfig(level=logging.DEBUG)

app = Flask(__name__)
CORS(app)

# ---------- CONEXÃO À BASE DE DADOS ----------
try:
    conn = psycopg2.connect(
        host="localhost",
        port="5433",
        dbname="AI4Governance",
        user="postgres",
        password="29344"
    )
except Exception as e:
    print("❌ Erro ao conectar à base de dados:", e)
    raise

INDEX_PATH = "faiss.index"
FAQ_EMBEDDINGS_PATH = "faq_embeddings.pkl"
embedding_model = SentenceTransformer('all-MiniLM-L12-v2')

# ---------- FUNÇÕES DE EMBEDDING E FAISS ----------
SAUDACOES = [
    "olá", "ola", "oi", "bom dia", "boa tarde", "boa noite", "oi", "hello", "hi"
]
SAUDACOES_PERGUNTAS = [
    "tudo bem", "como estás", "como está", "está tudo bem", "como vais", "tá bem"
]
RESPOSTA_SAUDACAO = "Olá! 👋 Como posso ajudar? Se tiver alguma dúvida, basta perguntar!"
RESPOSTA_TUDO_BEM = "Estou sempre pronto a ajudar! 😊 Em que posso ser útil?"

def detectar_saudacao(pergunta):
    texto = pergunta.strip().lower()
    palavras = set(texto.replace("?", "").replace("!", "").replace(".", "").split())
    for saud in SAUDACOES:
        for palavra in palavras:
            if saud == palavra:
                return RESPOSTA_SAUDACAO
    for p in SAUDACOES_PERGUNTAS:
        if p == texto:
            return RESPOSTA_TUDO_BEM
    return None

def preprocess_text(text):
    text = unidecode(text.lower())
    text = re.sub(r'[^\w\s]', '', text)
    stop_words = {'a', 'o', 'e', 'de', 'para', 'com', 'em', 'que', 'quem', 'como'}
    text = ' '.join(word for word in text.split() if word not in stop_words)
    return text

def get_faqs_from_db(chatbot_id=None):
    cur = conn.cursor()
    if chatbot_id:
        cur.execute("SELECT faq_id, pergunta, resposta, chatbot_id FROM FAQ WHERE chatbot_id = %s", (chatbot_id,))
    else:
        cur.execute("SELECT faq_id, pergunta, resposta, chatbot_id FROM FAQ")
    return cur.fetchall()

def build_faiss_index(chatbot_id=None):
    faqs = get_faqs_from_db(chatbot_id)
    perguntas = [preprocess_text(f[1]) for f in faqs]
    if not perguntas:
        emb_dim = 384
        embeddings = np.zeros((1, emb_dim), dtype=np.float32)
        index = faiss.IndexFlatIP(emb_dim)
    else:
        embeddings = embedding_model.encode(perguntas, show_progress_bar=True)
        embeddings = embeddings / np.linalg.norm(embeddings, axis=1, keepdims=True)
        index = faiss.IndexFlatIP(embeddings.shape[1])
        index.add(np.array(embeddings, dtype=np.float32))
    with open(FAQ_EMBEDDINGS_PATH, 'wb') as f:
        pickle.dump({'faqs': faqs, 'embeddings': embeddings}, f)
    faiss.write_index(index, INDEX_PATH)

def load_faiss_index():
    index = faiss.read_index(INDEX_PATH)
    with open(FAQ_EMBEDDINGS_PATH, 'rb') as f:
        data = pickle.load(f)
    return index, data['faqs'], data['embeddings']

if not (os.path.exists(INDEX_PATH) and os.path.exists(FAQ_EMBEDDINGS_PATH)):
    build_faiss_index()
faiss_index, faqs_db, faq_embeddings = load_faiss_index()

def pesquisar_faiss(pergunta, chatbot_id=None, k=1, min_sim=0.7):
    pergunta = preprocess_text(pergunta)
    if chatbot_id:
        faqs = [f for f in faqs_db if f[3] == int(chatbot_id)]
        if not faqs:
            return []
        perguntas = [preprocess_text(f[1]) for f in faqs]
        embeddings = embedding_model.encode(perguntas)
        embeddings = embeddings / np.linalg.norm(embeddings, axis=1, keepdims=True)
        index = faiss.IndexFlatIP(embeddings.shape[1])
        index.add(np.array(embeddings, dtype=np.float32))
        query_emb = embedding_model.encode([pergunta])
        query_emb = query_emb / np.linalg.norm(query_emb, axis=1, keepdims=True)
        D, I = index.search(np.array(query_emb, dtype=np.float32), min(k, len(faqs)))
        if D[0][0] < min_sim:
            return []
        results = []
        for idx in I[0]:
            faq_id, pergunta_faq, resposta_faq, chatbot_id_faq = faqs[idx]
            results.append({
                'faq_id': faq_id,
                'pergunta': pergunta_faq,
                'resposta': resposta_faq,
                'score': float(D[0][idx])
            })
        return results
    else:
        if len(faqs_db) == 0:
            return []
        query_emb = embedding_model.encode([pergunta])
        query_emb = query_emb / np.linalg.norm(query_emb, axis=1, keepdims=True)
        D, I = faiss_index.search(np.array(query_emb, dtype=np.float32), min(k, len(faqs_db)))
        if D[0][0] < min_sim:
            return []
        results = []
        for idx in I[0]:
            faq_id, pergunta_faq, resposta_faq, chatbot_id_faq = faqs_db[idx]
            results.append({
                'faq_id': faq_id,
                'pergunta': pergunta_faq,
                'resposta': resposta_faq,
                'score': float(D[0][idx])
            })
        return results

# ---------- SIMILARIDADE FUZZY ----------
def obter_faq_mais_semelhante(pergunta_utilizador, chatbot_id, threshold=60):
    cur = conn.cursor()
    cur.execute("SELECT pergunta, resposta FROM FAQ WHERE chatbot_id = %s", (chatbot_id,))
    faqs = cur.fetchall()
    pergunta_normalizada = pergunta_utilizador.strip().lower()
    melhor_pergunta = None
    melhor_resposta = None
    maior_score = 0
    for pergunta, resposta in faqs:
        pergunta_bd = pergunta.strip().lower()
        if pergunta_normalizada == pergunta_bd:
            return {"pergunta": pergunta, "resposta": resposta, "score": 100}
        score = fuzz.ratio(pergunta_normalizada, pergunta_bd)
        if score > maior_score:
            maior_score = score
            melhor_pergunta = pergunta
            melhor_resposta = resposta
    if maior_score >= threshold:
        return {"pergunta": melhor_pergunta, "resposta": melhor_resposta, "score": maior_score}
    else:
        return None

# ---------- CATEGORIAS ----------
@app.route("/chatbots/<int:chatbot_id>", methods=["DELETE"])
def eliminar_chatbot(chatbot_id):
    cur = conn.cursor()
    try:
        cur.execute("DELETE FROM FAQ_Relacionadas WHERE faq_id IN (SELECT faq_id FROM FAQ WHERE chatbot_id = %s)", (chatbot_id,))
        cur.execute("DELETE FROM FAQ_Documento WHERE faq_id IN (SELECT faq_id FROM FAQ WHERE chatbot_id = %s)", (chatbot_id,))
        cur.execute("DELETE FROM FAQ WHERE chatbot_id = %s", (chatbot_id,))
        cur.execute("DELETE FROM FonteResposta WHERE chatbot_id = %s", (chatbot_id,))
        cur.execute("DELETE FROM Chatbot WHERE chatbot_id = %s", (chatbot_id,))
        conn.commit()
        build_faiss_index()
        global faiss_index, faqs_db, faq_embeddings
        faiss_index, faqs_db, faq_embeddings = load_faiss_index()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/categorias", methods=["GET"])
def get_categorias():
    cur = conn.cursor()
    try:
        cur.execute("SELECT categoria_id, nome FROM Categoria")
        data = cur.fetchall()
        return jsonify([{"categoria_id": c[0], "nome": c[1]} for c in data])
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

# ---------- CHATBOTS ----------
@app.route("/chatbots", methods=["GET"])
def get_chatbots():
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT c.chatbot_id, c.nome, c.descricao, c.data_criacao, fr.fonte,
                   array_remove(array_agg(cc.categoria_id), NULL) as categorias
            FROM Chatbot c
            LEFT JOIN FonteResposta fr ON fr.chatbot_id = c.chatbot_id
            LEFT JOIN ChatbotCategoria cc ON cc.chatbot_id = c.chatbot_id
            GROUP BY c.chatbot_id, c.nome, c.descricao, c.data_criacao, fr.fonte
            ORDER BY c.chatbot_id ASC
        """)
        data = cur.fetchall()
        return jsonify([
            {
                "chatbot_id": row[0],
                "nome": row[1],
                "descricao": row[2],
                "data_criacao": row[3],
                "fonte": row[4] if row[4] else "faq",
                "categorias": row[5] if row[5] is not None else []
            }
            for row in data
        ])
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/chatbots", methods=["POST"])
def criar_chatbot():
    cur = conn.cursor()
    data = request.get_json()
    nome = data.get("nome", "").strip()
    descricao = data.get("descricao", "").strip()
    categorias = data.get("categorias", [])
    if not nome:
        return jsonify({"success": False, "error": "Nome obrigatório."}), 400
    try:
        cur.execute("SELECT chatbot_id FROM Chatbot WHERE LOWER(nome) = LOWER(%s)", (nome,))
        if cur.fetchone():
            return jsonify({"success": False, "error": "Já existe um chatbot com esse nome."}), 409

        cur.execute(
            "INSERT INTO Chatbot (nome, descricao) VALUES (%s, %s) RETURNING chatbot_id",
            (nome, descricao)
        )
        chatbot_id = cur.fetchone()[0]

        for categoria_id in categorias:
            cur.execute(
                "INSERT INTO ChatbotCategoria (chatbot_id, categoria_id) VALUES (%s, %s)",
                (chatbot_id, categoria_id)
            )

        cur.execute(
            "INSERT INTO FonteResposta (chatbot_id, fonte) VALUES (%s, %s)",
            (chatbot_id, "faq")
        )
        conn.commit()
        return jsonify({"success": True, "chatbot_id": chatbot_id})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/chatbots/<int:chatbot_id>", methods=["PUT"])
def atualizar_chatbot(chatbot_id):
    cur = conn.cursor()
    data = request.get_json()
    try:
        nome = data.get("nome", "").strip()
        descricao = data.get("descricao", "").strip()
        fonte = data.get("fonte", "faq")
        categorias = data.get("categorias", [])
        if not nome:
            return jsonify({"success": False, "error": "O nome do chatbot é obrigatório."}), 400

        cur.execute(
            "UPDATE Chatbot SET nome=%s, descricao=%s WHERE chatbot_id=%s",
            (nome, descricao, chatbot_id)
        )
        cur.execute("DELETE FROM ChatbotCategoria WHERE chatbot_id=%s", (chatbot_id,))
        for categoria_id in categorias:
            cur.execute(
                "INSERT INTO ChatbotCategoria (chatbot_id, categoria_id) VALUES (%s, %s)",
                (chatbot_id, categoria_id)
            )
        cur.execute("SELECT 1 FROM FonteResposta WHERE chatbot_id=%s", (chatbot_id,))
        if cur.fetchone():
            cur.execute("UPDATE FonteResposta SET fonte=%s WHERE chatbot_id=%s", (fonte, chatbot_id))
        else:
            cur.execute("INSERT INTO FonteResposta (chatbot_id, fonte) VALUES (%s, %s)", (chatbot_id, fonte))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        print(traceback.format_exc())
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/fonte/<int:chatbot_id>", methods=["GET"])
def obter_fonte_chatbot(chatbot_id):
    cur = conn.cursor()
    try:
        cur.execute("SELECT fonte FROM FonteResposta WHERE chatbot_id = %s", (chatbot_id,))
        row = cur.fetchone()
        if row:
            fonte = row[0] if row[0] else "faq"
            return jsonify({"success": True, "fonte": fonte})
        return jsonify({"success": False, "erro": "Chatbot não encontrado."}), 404
    except Exception as e:
        return jsonify({"success": False, "erro": str(e)}), 500

@app.route("/fonte", methods=["POST"])
def definir_fonte_chatbot():
    cur = conn.cursor()
    data = request.get_json()
    chatbot_id = data.get("chatbot_id")
    fonte = data.get("fonte")
    if fonte not in ["faq", "faiss", "faq+raga"]:
        return jsonify({"success": False, "erro": "Fonte inválida."}), 400
    try:
        cur.execute("SELECT 1 FROM FonteResposta WHERE chatbot_id = %s", (chatbot_id,))
        if not cur.fetchone():
            cur.execute("INSERT INTO FonteResposta (chatbot_id, fonte) VALUES (%s, %s)", (chatbot_id, fonte))
        else:
            cur.execute("UPDATE FonteResposta SET fonte = %s WHERE chatbot_id = %s", (fonte, chatbot_id))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "erro": str(e)}), 500

@app.route("/faq-categoria/<categoria>", methods=["GET"])
def obter_faq_por_categoria(categoria):
    cur = conn.cursor()
    try:
        chatbot_id = request.args.get("chatbot_id")
        if not chatbot_id:
            return jsonify({"success": False, "erro": "chatbot_id não fornecido."}), 400

        cur.execute("""
            SELECT f.faq_id, f.pergunta, f.resposta
            FROM FAQ f
            INNER JOIN Categoria c ON f.categoria_id = c.categoria_id
            WHERE LOWER(c.nome) = LOWER(%s) AND f.chatbot_id = %s
            ORDER BY RANDOM()
            LIMIT 1
        """, (categoria, chatbot_id))
        resultado = cur.fetchone()

        if resultado:
            return jsonify({
                "success": True,
                "faq_id": resultado[0],
                "pergunta": resultado[1],
                "resposta": resultado[2]
            })
        else:
            return jsonify({
                "success": False,
                "erro": f"Nenhuma FAQ encontrada para a categoria '{categoria}'."
            }), 404
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "erro": str(e)}), 500

@app.route("/faqs", methods=["GET"])
def get_faqs():
    cur = conn.cursor()
    try:
        cur.execute("SELECT faq_id, chatbot_id, designacao, pergunta, resposta FROM FAQ")
        data = cur.fetchall()
        return jsonify([
            {"faq_id": f[0], "chatbot_id": f[1], "designacao": f[2], "pergunta": f[3], "resposta": f[4]} for f in data
        ])
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/faqs/<int:faq_id>", methods=["GET"])
def get_faq_by_id(faq_id):
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT f.faq_id, f.chatbot_id, f.categoria_id, f.designacao, f.pergunta, f.resposta, f.idioma, f.links_documentos,
                   c.nome as categoria_nome, f.recomendado
            FROM FAQ f
            LEFT JOIN Categoria c ON f.categoria_id = c.categoria_id
            WHERE f.faq_id = %s
        """, (faq_id,))
        faq = cur.fetchone()
        if not faq:
            return jsonify({"success": False, "error": "FAQ não encontrada."}), 404

        return jsonify({
            "success": True,
            "faq": {
                "faq_id": faq[0],
                "chatbot_id": faq[1],
                "categoria_id": faq[2],
                "designacao": faq[3],
                "pergunta": faq[4],
                "resposta": faq[5],
                "idioma": faq[6],
                "links_documentos": faq[7],
                "categoria_nome": faq[8],
                "recomendado": faq[9]
            }
        })
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/faqs/chatbot/<int:chatbot_id>", methods=["GET"])
def get_faqs_por_chatbot(chatbot_id):
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT f.faq_id, c.nome, f.pergunta, f.resposta
            FROM FAQ f
            LEFT JOIN Categoria c ON f.categoria_id = c.categoria_id
            WHERE f.chatbot_id = %s
        """, (chatbot_id,))
        data = cur.fetchall()
        return jsonify([
            {
                "faq_id": row[0],
                "categoria": row[1],
                "pergunta": row[2],
                "resposta": row[3]
            }
            for row in data
        ])
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/faqs/<int:faq_id>", methods=["PUT"])
def update_faq(faq_id):
    cur = conn.cursor()
    data = request.get_json()
    try:
        pergunta = data.get("pergunta", "").strip()
        resposta = data.get("resposta", "").strip()
        idioma = data.get("idioma", "pt").strip()
        categorias = data.get("categorias", [])
        recomendado = data.get("recomendado", False)

        categoria_id = categorias[0] if categorias else None

        cur.execute("""
            UPDATE FAQ SET pergunta=%s, resposta=%s, idioma=%s, categoria_id=%s, recomendado=%s
            WHERE faq_id=%s
        """, (pergunta, resposta, idioma, categoria_id, recomendado, faq_id))

        conn.commit()
        build_faiss_index()
        global faiss_index, faqs_db, faq_embeddings
        faiss_index, faqs_db, faq_embeddings = load_faiss_index()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/faqs", methods=["POST"])
def add_faq():
    global faiss_index, faqs_db, faq_embeddings
    cur = conn.cursor()
    data = request.get_json()
    try:
        idioma = data.get("idioma", "").strip()
        if not idioma:
            return jsonify({"success": False, "error": "O campo 'idioma' é obrigatório."}), 400

        links_documentos = data.get("links_documentos", "").strip()
        recomendado = data.get("recomendado", False)

        cur.execute("""
            SELECT faq_id FROM FAQ
            WHERE chatbot_id = %s AND designacao = %s AND pergunta = %s AND resposta = %s AND idioma = %s
        """, (data["chatbot_id"], data["designacao"], data["pergunta"], data["resposta"], idioma))
        if cur.fetchone():
            return jsonify({"success": False, "error": "Esta FAQ já está inserida."}), 409

        cur.execute("""
            INSERT INTO FAQ (chatbot_id, categoria_id, designacao, pergunta, resposta, idioma, links_documentos, recomendado)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING faq_id
        """, (
            data["chatbot_id"],
            data["categoria_id"],
            data["designacao"],
            data["pergunta"],
            data["resposta"],
            idioma,
            links_documentos,
            recomendado
        ))
        faq_id = cur.fetchone()[0]

        if links_documentos:
            for link in links_documentos.split(','):
                link = link.strip()
                if link:
                    cur.execute(
                        "INSERT INTO FAQ_Documento (faq_id, link) VALUES (%s, %s)",
                        (faq_id, link)
                    )

        if "relacionadas" in data and data["relacionadas"].strip():
            for rel_id in data["relacionadas"].split(','):
                cur.execute("INSERT INTO FAQ_Relacionadas (faq_id, faq_relacionada_id) VALUES (%s, %s)", (faq_id, int(rel_id.strip())))
        conn.commit()
        build_faiss_index()
        faiss_index, faqs_db, faq_embeddings = load_faiss_index()
        return jsonify({"success": True, "faq_id": faq_id})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/faqs/<int:faq_id>", methods=["DELETE"])
def delete_faq(faq_id):
    cur = conn.cursor()
    try:
        cur.execute("DELETE FROM FAQ WHERE faq_id = %s", (faq_id,))
        conn.commit()
        build_faiss_index()
        global faiss_index, faqs_db, faq_embeddings
        faiss_index, faqs_db, faq_embeddings = load_faiss_index()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/upload-faq-docx", methods=["POST"])
def upload_faq_docx():
    cur = conn.cursor()
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "Ficheiro não enviado."}), 400

    file = request.files['file']
    chatbot_id_raw = request.form.get("chatbot_id")
    if not chatbot_id_raw:
        return jsonify({"success": False, "error": "Chatbot ID não fornecido."}), 400

    def normalizar_idioma(valor):
        if not valor:
            return "pt"
        valor = valor.strip().lower()
        if valor.startswith("port"):
            return "pt"
        if valor.startswith("ingl"):
            return "en"
        return valor[:2]

    try:
        doc = docx.Document(file)
        dados = {}

        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    chave_raw = row.cells[0].text.strip().lower().replace("’", "'")
                    valor = row.cells[1].text.strip()
                    chave = chave_raw.replace(":", "").strip()
                    if chave and valor:
                        dados[chave] = valor

        if not dados.get("designação da faq") or not dados.get("questão") or not dados.get("resposta"):
            raise Exception("Faltam campos obrigatórios: designação, questão ou resposta.")

        designacao = dados.get("designação da faq")
        pergunta = dados.get("questão")
        resposta = dados.get("resposta")
        categoria = dados.get("categoria")
        idioma_lido = dados.get("idioma", "Português")
        idioma = normalizar_idioma(idioma_lido)

        links_documentos = ""
        for key in ["documentos associados", "links de documentos"]:
            if key in dados:
                links_documentos = dados[key]
                break

        chatbot_ids = []
        if chatbot_id_raw == "todos":
            cur.execute("SELECT chatbot_id FROM Chatbot")
            chatbot_ids = [row[0] for row in cur.fetchall()]
        else:
            chatbot_ids = [int(chatbot_id_raw)]

        for chatbot_id in chatbot_ids:
            cur.execute("""
                SELECT faq_id FROM FAQ
                WHERE chatbot_id = %s AND designacao = %s AND pergunta = %s AND resposta = %s AND idioma = %s
            """, (chatbot_id, designacao, pergunta, resposta, idioma))
            if cur.fetchone():
                continue

            cur.execute("""
                INSERT INTO FAQ (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos)
                VALUES (%s, %s, %s, %s, %s, %s)
                RETURNING faq_id
            """, (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos))
            faq_id = cur.fetchone()[0]

            if categoria:
                cur.execute("SELECT categoria_id FROM Categoria WHERE nome ILIKE %s", (categoria,))
                result = cur.fetchone()
                if result:
                    cur.execute("UPDATE FAQ SET categoria_id = %s WHERE faq_id = %s", (result[0], faq_id))

            if links_documentos:
                for link in links_documentos.split(','):
                    link = link.strip()
                    if link:
                        cur.execute(
                            "INSERT INTO FAQ_Documento (faq_id, link) VALUES (%s, %s)",
                            (faq_id, link)
                        )
        conn.commit()
        build_faiss_index()
        global faiss_index, faqs_db, faq_embeddings
        faiss_index, faqs_db, faq_embeddings = load_faiss_index()
        return jsonify({"success": True, "message": "FAQ e links inseridos com sucesso."})

    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/upload-faq-docx-multiplos", methods=["POST"])
def upload_faq_docx_multiplos():
    cur = conn.cursor()
    if 'files' not in request.files:
        return jsonify({"success": False, "error": "Ficheiros não enviados."}), 400

    chatbot_id_raw = request.form.get("chatbot_id")
    if not chatbot_id_raw:
        return jsonify({"success": False, "error": "Chatbot ID não fornecido."}), 400

    files = request.files.getlist('files')
    total_inseridas = 0
    erros = []

    def normalizar_idioma(valor):
        if not valor:
            return "pt"
        valor = valor.strip().lower()
        if valor.startswith("port"):
            return "pt"
        if valor.startswith("ingl"):
            return "en"
        return valor[:2]

    for file in files:
        try:
            doc = docx.Document(file)
            dados = {}
            for table in doc.tables:
                for row in table.rows:
                    if len(row.cells) >= 2:
                        chave_raw = row.cells[0].text.strip().lower().replace("’", "'")
                        valor = row.cells[1].text.strip()
                        chave = chave_raw.replace(":", "").strip()
                        if chave and valor:
                            dados[chave] = valor
            if not dados.get("designação da faq") or not dados.get("questão") or not dados.get("resposta"):
                raise Exception("Faltam campos obrigatórios: designação, questão ou resposta.")

            designacao = dados.get("designação da faq")
            pergunta = dados.get("questão")
            resposta = dados.get("resposta")
            categoria = dados.get("categoria")
            idioma_lido = dados.get("idioma", "Português")
            idioma = normalizar_idioma(idioma_lido)

            links_documentos = ""
            for key in ["documentos associados", "links de documentos"]:
                if key in dados:
                    links_documentos = dados[key]
                    break

            chatbot_ids = []
            if chatbot_id_raw == "todos":
                cur.execute("SELECT chatbot_id FROM Chatbot")
                chatbot_ids = [row[0] for row in cur.fetchall()]
            else:
                chatbot_ids = [int(chatbot_id_raw)]

            for chatbot_id in chatbot_ids:
                cur.execute("""
                    SELECT faq_id FROM FAQ
                    WHERE chatbot_id = %s AND designacao = %s AND pergunta = %s AND resposta = %s AND idioma = %s
                """, (chatbot_id, designacao, pergunta, resposta, idioma))
                if cur.fetchone():
                    continue

                cur.execute("""
                    INSERT INTO FAQ (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    RETURNING faq_id
                """, (chatbot_id, designacao, pergunta, resposta, idioma, links_documentos))
                faq_id = cur.fetchone()[0]

                if categoria:
                    cur.execute("SELECT categoria_id FROM Categoria WHERE nome ILIKE %s", (categoria,))
                    result = cur.fetchone()
                    if result:
                        cur.execute("UPDATE FAQ SET categoria_id = %s WHERE faq_id = %s", (result[0], faq_id))

                if links_documentos:
                    for link in links_documentos.split(','):
                        link = link.strip()
                        if link:
                            cur.execute(
                                "INSERT INTO FAQ_Documento (faq_id, link) VALUES (%s, %s)",
                                (faq_id, link)
                            )
                total_inseridas += 1
        except Exception as e:
            erros.append(str(e))
            conn.rollback()
    conn.commit()
    build_faiss_index()
    global faiss_index, faqs_db, faq_embeddings
    faiss_index, faqs_db, faq_embeddings = load_faiss_index()
    return jsonify({"success": True, "inseridas": total_inseridas, "erros": erros})

@app.route("/faqs/detalhes", methods=["GET"])
def get_faqs_detalhes():
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT f.faq_id, f.chatbot_id, f.designacao, f.pergunta, f.resposta, f.idioma, f.links_documentos,
                   f.categoria_id, c.nome AS categoria_nome, ch.nome AS chatbot_nome, f.recomendado
            FROM FAQ f
            LEFT JOIN Categoria c ON f.categoria_id = c.categoria_id
            LEFT JOIN Chatbot ch ON f.chatbot_id = ch.chatbot_id
            ORDER BY f.faq_id
        """)
        data = cur.fetchall()
        return jsonify([
            {
                "faq_id": r[0],
                "chatbot_id": r[1],
                "designacao": r[2],
                "pergunta": r[3],
                "resposta": r[4],
                "idioma": r[5],
                "links_documentos": r[6],
                "categoria_id": r[7],
                "categoria_nome": r[8],
                "chatbot_nome": r[9],
                "recomendado": r[10]
            }
            for r in data
        ])
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/obter-resposta", methods=["POST"])
def obter_resposta():
    cur = conn.cursor()
    dados = request.get_json()
    pergunta = dados.get("pergunta", "").strip()
    chatbot_id = dados.get("chatbot_id")
    fonte = dados.get("fonte", "faq")

    saudacao = detectar_saudacao(pergunta)
    if saudacao:
        return jsonify({
            "success": True,
            "fonte": "SAUDACAO",
            "resposta": saudacao,
            "faq_id": None,
            "categoria_id": None
        })

    if not pergunta or (len(pergunta) < 4 and not any(char.isalpha() for char in pergunta)):
        return jsonify({
            "success": False,
            "erro": "Pergunta demasiado curta ou não reconhecida como válida."
        })

    try:
        if fonte == "faq":
            resultado = obter_faq_mais_semelhante(pergunta, chatbot_id)
            if resultado:
                cur.execute("""
                    SELECT faq_id, categoria_id FROM FAQ
                    WHERE LOWER(pergunta) = LOWER(%s) AND chatbot_id = %s
                """, (resultado["pergunta"], chatbot_id))
                row = cur.fetchone()
                faq_id, categoria_id = row if row else (None, None)
                return jsonify({
                    "success": True,
                    "fonte": "FAQ",
                    "resposta": resultado["resposta"],
                    "faq_id": faq_id,
                    "categoria_id": categoria_id,
                    "score": resultado["score"]
                })
            return jsonify({"success": False, "erro": "Pergunta não encontrada nas FAQs."})

        elif fonte == "faiss":
            faiss_resultados = pesquisar_faiss(pergunta, chatbot_id=chatbot_id, k=1, min_sim=0.7)
            if faiss_resultados:
                return jsonify({
                    "success": True,
                    "fonte": "FAISS",
                    "resposta": faiss_resultados[0]['resposta'],
                    "faq_id": faiss_resultados[0]['faq_id'],
                    "score": faiss_resultados[0]['score']
                })
            else:
                resultado = obter_faq_mais_semelhante(pergunta, chatbot_id, threshold=80)
                if resultado:
                    cur.execute("""
                        SELECT faq_id, categoria_id FROM FAQ
                        WHERE LOWER(pergunta) = LOWER(%s) AND chatbot_id = %s
                    """, (resultado["pergunta"], chatbot_id))
                    row = cur.fetchone()
                    faq_id, categoria_id = row if row else (None, None)
                    return jsonify({
                        "success": True,
                        "fonte": "FUZZY",
                        "resposta": resultado["resposta"],
                        "faq_id": faq_id,
                        "categoria_id": categoria_id,
                        "score": resultado["score"]
                    })
                return jsonify({
                    "success": False,
                    "erro": "Não encontrei nenhuma resposta suficientemente semelhante na base de dados."
                })

        elif fonte == "faq+raga":
            resultado = obter_faq_mais_semelhante(pergunta, chatbot_id)
            if resultado:
                cur.execute("""
                    SELECT faq_id, categoria_id FROM FAQ
                    WHERE LOWER(pergunta) = LOWER(%s) AND chatbot_id = %s
                """, (resultado["pergunta"], chatbot_id))
                row = cur.fetchone()
                faq_id, categoria_id = row if row else (None, None)
                return jsonify({
                    "success": True,
                    "fonte": "FAQ",
                    "resposta": resultado["resposta"],
                    "faq_id": faq_id,
                    "categoria_id": categoria_id,
                    "score": resultado["score"]
                })
            else:
                faiss_resultados = pesquisar_faiss(pergunta, chatbot_id=chatbot_id, k=1, min_sim=0.7)
                if faiss_resultados:
                    return jsonify({
                        "success": True,
                        "fonte": "RAG",
                        "resposta": faiss_resultados[0]['resposta'],
                        "faq_id": faiss_resultados[0]['faq_id'],
                        "score": faiss_resultados[0]['score']
                    })
                else:
                    return jsonify({
                        "success": False,
                        "erro": "Não encontrei nenhuma resposta suficientemente semelhante na base de dados (RAG/FAISS)."
                    })
        else:
            return jsonify({"success": False, "erro": "Fonte inválida."}), 400
    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"success": False, "erro": str(e)}), 500

@app.route("/perguntas-semelhantes", methods=["POST"])
def perguntas_semelhantes():
    cur = conn.cursor()
    dados = request.get_json()
    pergunta_atual = dados.get("pergunta", "")
    chatbot_id = dados.get("chatbot_id")
    try:
        cur.execute("""
            SELECT f.categoria_id
            FROM FAQ f
            WHERE LOWER(f.pergunta) = LOWER(%s) AND f.chatbot_id = %s
        """, (pergunta_atual.strip().lower(), chatbot_id))
        categoria_row = cur.fetchone()
        if not categoria_row:
            return jsonify({"success": True, "sugestoes": []})
        categoria_id = categoria_row[0]
        cur.execute("""
            SELECT pergunta
            FROM FAQ
            WHERE categoria_id = %s AND LOWER(pergunta) != LOWER(%s) AND chatbot_id = %s
            ORDER BY RANDOM()
            LIMIT 2
        """, (categoria_id, pergunta_atual.strip().lower(), chatbot_id))
        sugestoes = [row[0] for row in cur.fetchall()]
        return jsonify({"success": True, "sugestoes": sugestoes})
    except Exception as e:
        return jsonify({"success": False, "erro": str(e)}), 500

@app.route("/chatbot/<int:chatbot_id>", methods=["GET"])
def obter_nome_chatbot(chatbot_id):
    cur = conn.cursor()
    try:
        cur.execute("SELECT nome FROM Chatbot WHERE chatbot_id = %s", (chatbot_id,))
        row = cur.fetchone()
        if row:
            return jsonify({"success": True, "nome": row[0]})
        return jsonify({"success": False, "erro": "Chatbot não encontrado."}), 404
    except Exception as e:
        return jsonify({"success": False, "erro": str(e)}), 500

@app.route("/rebuild-faiss", methods=["POST"])
def rebuild_faiss():
    build_faiss_index()
    global faiss_index, faqs_db, faq_embeddings
    faiss_index, faqs_db, faq_embeddings = load_faiss_index()
    return jsonify({"success": True, "msg": "FAISS index rebuilt."})

# ---------- INICIAR SERVIDOR ----------
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)